# mainKDR/views.py - ИСПРАВЛЕННАЯ ВЕРСИЯ
from django.shortcuts import render, redirect
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.contrib.auth import login, authenticate, logout
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, Http404
from django.conf import settings
from io import BytesIO
import zipfile
import os
import re
from datetime import datetime

# Словари для преобразования месяцев
MONTHS_GENITIVE = {
    1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
    5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
    9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
}

MONTHS_NOMINATIVE = {
    1: 'январь', 2: 'февраль', 3: 'март', 4: 'апрель',
    5: 'май', 6: 'июнь', 7: 'июль', 8: 'август',
    9: 'сентябрь', 10: 'октябрь', 11: 'ноябрь', 12: 'декабрь'
}

MONTH_NAME_TO_NUMBER = {
    'января': 1, 'февраля': 2, 'марта': 3, 'апреля': 4,
    'мая': 5, 'июня': 6, 'июля': 7, 'августа': 8,
    'сентября': 9, 'октября': 10, 'ноября': 11, 'декабря': 12
}


def decline_fio_simple(full_name):
    """
    Простое склонение ФИО в родительный падеж (кого?)
    """
    if not full_name or full_name.strip() == '':
        return full_name
    
    parts = full_name.strip().split()
    
    if len(parts) == 1:
        return full_name
    
    surname = parts[0]
    name = parts[1] if len(parts) > 1 else ''
    patronymic = parts[2] if len(parts) > 2 else ''
    
    # Склоняем фамилию
    if surname.endswith('ов') or surname.endswith('ев'):
        surname = surname[:-2] + 'ова'
    elif surname.endswith('ёв'):
        surname = surname[:-2] + 'ёва'
    elif surname.endswith('ин') or surname.endswith('ын'):
        surname = surname[:-2] + 'ина'
    elif surname.endswith('ий'):
        surname = surname[:-2] + 'его'
    elif surname.endswith('ой') or surname.endswith('ый'):
        surname = surname[:-2] + 'ого'
    elif surname.endswith('ая'):
        surname = surname[:-2] + 'ой'
    elif surname.endswith('ан'):
        surname = surname[:-2] + 'ана'
    elif surname.endswith('ян'):
        surname = surname[:-2] + 'яна'
    elif surname.endswith('ман'):
        surname = surname[:-2] + 'мана'
    elif surname.endswith('цман'):
        surname = surname[:-4] + 'цмана'
    elif surname.endswith('ко') or surname.endswith('ак') or surname.endswith('ых') or surname.endswith('их'):
        pass
    elif surname.endswith('а'):
        surname = surname[:-1] + 'ы'
    elif surname.endswith('я'):
        surname = surname[:-1] + 'и'
    
    # Склоняем имя
    if name:
        if name.endswith('й'):
            name = name[:-1] + 'я'
        elif name.endswith('а'):
            name = name[:-1] + 'ы'
        elif name.endswith('я'):
            name = name[:-1] + 'и'
        elif name.endswith('ь'):
            name = name[:-1] + 'я'
        else:
            name = name + 'а'
    
    # Склоняем отчество
    if patronymic:
        if patronymic.endswith('на'):
            patronymic = patronymic[:-2] + 'ны'
        elif patronymic.endswith('чна') or patronymic.endswith('вна'):
            patronymic = patronymic[:-2] + 'чны' if patronymic.endswith('чна') else patronymic[:-2] + 'вны'
        elif patronymic.endswith('ч'):
            patronymic = patronymic + 'а'
        elif patronymic.endswith('вна'):
            patronymic = patronymic[:-2] + 'вны'
        else:
            patronymic = patronymic + 'а'
    
    if patronymic:
        return f"{surname} {name} {patronymic}"
    elif name:
        return f"{surname} {name}"
    else:
        return surname


def decline_vid(vid):
    """Склоняет вид практики в винительный падеж"""
    if not vid:
        return ''
    
    vid_lower = vid.lower()
    
    if vid_lower == 'учебная':
        return 'учебную'
    elif vid_lower == 'производственная':
        return 'производственную'
    elif vid_lower == 'преддипломная':
        return 'преддипломную'
    else:
        return vid


def parse_date_to_components(date_str):
    """
    Преобразует дату из формата YYYY-MM-DD в компоненты:
    день (число), месяц в родительном падеже, год
    Возвращает tuple (day, month_text, year)
    """
    if not date_str:
        return ('', '', '')
    
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        day = str(date_obj.day)
        month_text = MONTHS_GENITIVE.get(date_obj.month, '')
        year = str(date_obj.year)
        return (day, month_text, year)
    except:
        return ('', '', '')


def format_date_for_calendar(day, month_text, year):
    """
    Преобразует компоненты даты (день, месяц текстом, год) в формат YYYY-MM-DD для календаря
    """
    if not day or not month_text or not year:
        return ''
    
    month_num = MONTH_NAME_TO_NUMBER.get(month_text.lower(), 1)
    try:
        return f"{year}-{int(month_num):02d}-{int(day):02d}"
    except:
        return ''


try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("WARNING: python-docx not installed. Run: pip install python-docx")

from .models import DocWork, AttestList, Students, Contracts, ControlList, TypePractic, Appendix


def register(request):
    if request.method == "POST":
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            username = form.cleaned_data.get('username')
            
            Students.objects.create(user=user, familia='', name='', otchestvo='')
            DocWork.objects.create(user=user)
            AttestList.objects.create(user=user)
            Contracts.objects.create(user=user)
            ControlList.objects.create(user=user)
            TypePractic.objects.create(user=user)
            Appendix.objects.create(user=user)
            
            messages.success(request, f'Аккаунт {username} успешно создан!')
            return redirect('/login/')
        else:
            for error in form.errors.values():
                messages.error(request, error)
    else:
        form = UserCreationForm()
    
    return render(request, "register.html", {"form": form})


def loginf(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        
        if user is not None:
            login(request, user)
            messages.success(request, f'Добро пожаловать, {username}!')
            return redirect('/profile/')
        else:
            messages.error(request, 'Неверное имя пользователя или пароль.')
    
    return render(request, 'login.html', {'form': AuthenticationForm()})


def replace_variables_in_doc(doc, variables):
    """Замена переменных в документе"""
    
    def replace_text(text, variables):
        if not text:
            return text
        
        for pattern, replacement in variables.items():
            if replacement is None:
                replacement = ''
            
            str_replacement = str(replacement)
            
            formats = [
                pattern,
                f'[[{pattern}]]',
                f'[[{pattern}]]{{.underline}}',
                f'[[{pattern}]]{{.underline}}]',
                f'[{pattern}]',
                f'[{pattern}]{{.underline}}',
                f'[{pattern}]{{.underline}}]',
                f'[[{pattern.strip("{}")}]]',
                f'[{pattern.strip("{}")}]',
            ]
            
            for fmt in formats:
                if fmt in text:
                    text = text.replace(fmt, str_replacement)
        
        return text
    
    for paragraph in doc.paragraphs:
        original = paragraph.text
        if original:
            new_text = replace_text(original, variables)
            if new_text != original:
                paragraph.text = new_text
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original = paragraph.text
                    if original:
                        new_text = replace_text(original, variables)
                        if new_text != original:
                            paragraph.text = new_text
    
    return doc


def generate_document(template_path, variables):
    """Генерация документа"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document(template_path)
        doc = replace_variables_in_doc(doc, variables)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        print(f"Ошибка: {e}")
        return None


@login_required(login_url='/login/')
def profile(request):
    current_user = request.user
    
    student, _ = Students.objects.get_or_create(user=current_user)
    practice, _ = DocWork.objects.get_or_create(user=current_user)
    attest, _ = AttestList.objects.get_or_create(user=current_user)
    contract, _ = Contracts.objects.get_or_create(user=current_user)
    control, _ = ControlList.objects.get_or_create(user=current_user)
    type_practic, _ = TypePractic.objects.get_or_create(user=current_user)
    appendix, _ = Appendix.objects.get_or_create(user=current_user)
    
    if request.method == 'POST':
        if 'generate_docs' in request.POST:
            if not DOCX_AVAILABLE:
                messages.error(request, '❌ Установите python-docx')
                return redirect('/profile/')
            
            # Получаем значения из POST
            fio_value = request.POST.get('attest_fio', '') or f"{request.POST.get('familia', '')} {request.POST.get('name', '')} {request.POST.get('otchestvo', '')}".strip()
            spec_value = request.POST.get('spec', '')
            grupa_value = request.POST.get('grupa', '')
            kurs_value = request.POST.get('kurs', '')
            obuch_value = request.POST.get('obuch', 'очная')
            
            # === ОСНОВНАЯ ЛОГИКА ПОЛУЧЕНИЯ ДАТ ===
            calendar_start = request.POST.get('data_start', '')
            calendar_end = request.POST.get('data_end', '')
            
            text_start_day = request.POST.get('start', '')
            text_start_month = request.POST.get('stmo', '')
            text_start_year = request.POST.get('year', '')
            text_end_day = request.POST.get('final', '')
            text_end_month = request.POST.get('fnmo', '')
            text_end_year = request.POST.get('year', '')
            
            cal_start_day, cal_start_month, cal_start_year = parse_date_to_components(calendar_start)
            cal_end_day, cal_end_month, cal_end_year = parse_date_to_components(calendar_end)
            
            data_day = text_start_day if text_start_day else cal_start_day
            data_month = text_start_month if text_start_month else cal_start_month
            data_year = text_start_year if text_start_year else cal_start_year
            
            data3_day = text_end_day if text_end_day else cal_end_day
            data4_month = text_end_month if text_end_month else cal_end_month
            god1_year = text_end_year if text_end_year else cal_end_year
            
            if not god1_year:
                god1_year = data_year
            
            if not data_month:
                data_month = 'мая'
            if not data4_month:
                data4_month = 'мая'
            
            vid_value = request.POST.get('vid', '')
            kod_value = request.POST.get('kod', '')
            mesto_value = request.POST.get('mesto', '')
            adress_value = request.POST.get('attest_adress', '')
            ruka_value = request.POST.get('ruka', '')
            
            variables = {
                '{{fio}}': decline_fio_simple(fio_value),
                '{{spec}}': spec_value,
                '{{grupa}}': grupa_value,
                '{{kurs}}': kurs_value,
                '{{obuch}}': obuch_value,
                '{{data}}': data_day,
                '{{data2}}': data_month,
                '{{god}}': data_year,
                '{{data3}}': data3_day,
                '{{data4}}': data4_month,
                '{{god1}}': god1_year,
                '{{vid}}': decline_vid(vid_value),
                '{{kod}}': kod_value,
                '{{mesto}}': mesto_value,
                '{{adress}}': adress_value,
                '{{ruka}}': ruka_value,
                '{{data_start}}': data_day,
                '{{data2_start}}': data_month,
                '{{god_start}}': data_year,
                '{{data_end}}': data3_day,
                '{{data4_end}}': data4_month,
                '{{god_end}}': god1_year,
                '{{profORG}}': request.POST.get('profORG', ''),
                '{{ryco}}': request.POST.get('ryco', ''),
                '{{smart}}': request.POST.get('smart', ''),
                '{{rec}}': request.POST.get('rec', ''),
                '{{gru}}': request.POST.get('control_gru', ''),
                '{{code}}': request.POST.get('control_code', ''),
                '{{napra}}': request.POST.get('napra', ''),
                '{{rycSPO}}': request.POST.get('rycSPO', ''),
                '{{date}}': request.POST.get('date_check', ''),
                '{{pred}}': request.POST.get('pred', ''),
                '{{rycP}}': request.POST.get('rycP', ''),
                '{{fioO}}': request.POST.get('fioO', ''),
                '{{pri}}': request.POST.get('pri', ''),
                '{{modul}}': request.POST.get('modul', ''),
                '{{start}}': data_day,
                '{{stmo}}': data_month,
                '{{year}}': data_year,
                '{{final}}': data3_day,
                '{{fnmo}}': data4_month,
                '{{akd}}': request.POST.get('akd', ''),
                '{{name}}': request.POST.get('name_spec', ''),
                '{{date2}}': request.POST.get('date_birth2', ''),
                '{{fio2}}': request.POST.get('fio2', ''),
                '{{fioRuk}}': request.POST.get('fioRuk', ''),
                '{{num}}': request.POST.get('num', ''),
                '{{fioProfRuk}}': request.POST.get('fioProfRuk', ''),
                '{{doljnost}}': request.POST.get('doljnost', ''),
                '{{numProfRuk}}': request.POST.get('numProfRuk', ''),
                '{{org}}': request.POST.get('org', ''),
                '{{ruk}}': request.POST.get('ruk', ''),
                '{{fioOtv}}': request.POST.get('fioOtv', ''),
                '{{doljOtv}}': request.POST.get('doljOtv', ''),
                '{{numberOtv}}': request.POST.get('numberOtv', ''),
                '{{namePomech}}': request.POST.get('namePomech', ''),
                '{{addres}}': request.POST.get('appendix_addres', ''),
            }
            
            templates_dir = os.path.join(settings.BASE_DIR, 'templates', 'docs')
            
            templates = {
                'Аттестационный_лист.docx': os.path.join(templates_dir, 'шаблон АТТЕСТАЦИОННЫЙ ЛИСТ.docx'),
                'Договор.docx': os.path.join(templates_dir, 'шаблон ДОГОВОР.docx'),
                'Лист_контроля.docx': os.path.join(templates_dir, 'шаблон Лист контроля.docx'),
                'Приложение_1.docx': os.path.join(templates_dir, 'шаблон ПРИЛОЖЕНИЕ №1.docx'),
                'Приложение_2.docx': os.path.join(templates_dir, 'шаблон ПРИЛОЖЕНИЕ №2.docx'),
            }
            
            zip_buffer = BytesIO()
            generated = 0
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for name, path in templates.items():
                    if os.path.exists(path):
                        try:
                            doc = Document(path)
                            
                            for p in doc.paragraphs:
                                original_text = p.text
                                new_text = original_text
                                for key, val in variables.items():
                                    if val and key in new_text:
                                        new_text = new_text.replace(key, str(val))
                                if new_text != original_text:
                                    p.text = new_text
                            
                            for table in doc.tables:
                                for row in table.rows:
                                    for cell in row.cells:
                                        for p in cell.paragraphs:
                                            original_text = p.text
                                            new_text = original_text
                                            for key, val in variables.items():
                                                if val and key in new_text:
                                                    new_text = new_text.replace(key, str(val))
                                            if new_text != original_text:
                                                p.text = new_text
                            
                            buffer = BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            zip_file.writestr(name, buffer.getvalue())
                            generated += 1
                            print(f"✓ {name}")
                        except Exception as e:
                            print(f"✗ {name}: {e}")
            
            if generated == 0:
                messages.error(request, '❌ Ошибка генерации')
                return redirect('/profile/')
            
            zip_buffer.seek(0)
            response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename="Документы_практики.zip"'
            messages.success(request, f'✅ Сгенерировано {generated} документов!')
            return response
        
        else:  # save_data
            student.familia = request.POST.get('familia', '')
            student.name = request.POST.get('name', '')
            student.otchestvo = request.POST.get('otchestvo', '')
            student.save()
            
            date_birth_str = request.POST.get('date_birth', '')
            birth_day, birth_month_text, birth_year = parse_date_to_components(date_birth_str)
            formatted_birth_date = f"{birth_day} {birth_month_text} {birth_year}" if birth_day else ''
            
            data_start_str = request.POST.get('data_start', '')
            start_day, start_month_text, start_year = parse_date_to_components(data_start_str)
            
            data_end_str = request.POST.get('data_end', '')
            end_day, end_month_text, end_year = parse_date_to_components(data_end_str)
            
            if not start_day:
                start_day = request.POST.get('start', '')
            if not start_month_text:
                start_month_text = request.POST.get('stmo', '')
            if not start_year:
                start_year = request.POST.get('year', '')
            
            if not end_day:
                end_day = request.POST.get('final', '')
            if not end_month_text:
                end_month_text = request.POST.get('fnmo', '')
            if not end_year:
                end_year = request.POST.get('year', '')
            
            practice.modul = request.POST.get('modul', '')
            practice.start = start_day
            practice.stmo = start_month_text
            practice.year = start_year
            practice.final = end_day
            practice.fnmo = end_month_text
            practice.kod = request.POST.get('kod', '')
            practice.name_spec = request.POST.get('name_spec', '')
            practice.akd = request.POST.get('akd', '') or None
            practice.fio = request.POST.get('fio', '')
            practice.date_birth = formatted_birth_date
            practice.fio2 = request.POST.get('fio2', '')
            practice.date_birth2 = request.POST.get('date_birth2', '')
            practice.fioRuk = request.POST.get('fioRuk', '')
            practice.num = request.POST.get('num', '')
            practice.fioProfRuk = request.POST.get('fioProfRuk', '')
            practice.doljnost = request.POST.get('doljnost', '')
            practice.numProfRuk = request.POST.get('numProfRuk', '')
            practice.org = request.POST.get('org', '')
            practice.ruk = request.POST.get('ruk', '')
            practice.fioOtv = request.POST.get('fioOtv', '')
            practice.doljOtv = request.POST.get('doljOtv', '')
            practice.numberOtv = request.POST.get('numberOtv', '')
            practice.save()
            
            attest.fio = request.POST.get('attest_fio', '')
            attest.spec = request.POST.get('spec', '')
            attest.grupa = request.POST.get('grupa', '')
            attest.kurs = request.POST.get('kurs', '')
            attest.obuch = request.POST.get('obuch', '')
            attest.data_start = start_day
            attest.data2_start = start_month_text
            attest.god_start = start_year
            attest.data_end = end_day
            attest.data4_end = end_month_text
            attest.god_end = end_year
            attest.vid = request.POST.get('vid', '')
            attest.kod = request.POST.get('attest_kod', '')
            attest.mesto = request.POST.get('mesto', '')
            attest.adress = request.POST.get('attest_adress', '')
            attest.ruka = request.POST.get('ruka', '')
            attest.save()
            
            contract.num_dogovor = request.POST.get('num_dogovor', '')
            contract.profORG = request.POST.get('profORG', '')
            contract.ryco = request.POST.get('ryco', '')
            contract.adress = request.POST.get('contract_adress', '')
            contract.smart = request.POST.get('smart', '')
            contract.rec = request.POST.get('rec', '')
            contract.save()
            
            control.gru = request.POST.get('control_gru', '')
            control.code = request.POST.get('control_code', '')
            control.napra = request.POST.get('napra', '')
            control.data = request.POST.get('control_data', '')
            control.rycSPO = request.POST.get('rycSPO', '')
            control.date_check = request.POST.get('date_check', '')
            control.pred = request.POST.get('pred', '')
            control.rycP = request.POST.get('rycP', '')
            control.fio = request.POST.get('control_fio', '')
            control.fioO = request.POST.get('fioO', '')
            control.pri = request.POST.get('pri', '')
            control.save()
            
            type_practic.field1 = request.POST.get('field1', '')
            type_practic.save()
            
            appendix.addres = request.POST.get('appendix_addres', '')
            appendix.namePomech = request.POST.get('namePomech', '')
            appendix.save()
            
            messages.success(request, '✅ Данные сохранены!')
            return redirect('/profile/')
    
    # GET запрос
    start_date_for_calendar = format_date_for_calendar(practice.start, practice.stmo, practice.year)
    end_date_for_calendar = format_date_for_calendar(practice.final, practice.fnmo, practice.year)
    
    birth_date_for_calendar = ''
    if practice.date_birth:
        parts = practice.date_birth.split()
        if len(parts) >= 3:
            birth_date_for_calendar = format_date_for_calendar(parts[0], parts[1], parts[2])
    
    context = {
        'username': current_user.username,
        'student_familia': student.familia,
        'student_name': student.name,
        'student_otchestvo': student.otchestvo,
        'modul': practice.modul,
        'start': practice.start,
        'stmo': practice.stmo,
        'year': practice.year,
        'final': practice.final,
        'fnmo': practice.fnmo,
        'kod': practice.kod,
        'name_spec': practice.name_spec,
        'akd': practice.akd,
        'fio': practice.fio,
        'date_birth': birth_date_for_calendar,
        'fio2': practice.fio2,
        'date_birth2': practice.date_birth2,
        'fioRuk': practice.fioRuk,
        'num': practice.num,
        'fioProfRuk': practice.fioProfRuk,
        'doljnost': practice.doljnost,
        'numProfRuk': practice.numProfRuk,
        'org': practice.org,
        'ruk': practice.ruk,
        'fioOtv': practice.fioOtv,
        'doljOtv': practice.doljOtv,
        'numberOtv': practice.numberOtv,
        'attest_fio': attest.fio,
        'spec': attest.spec,
        'grupa': attest.grupa,
        'kurs': attest.kurs,
        'obuch': attest.obuch,
        'data_start': start_date_for_calendar,
        'data2_start': attest.data2_start,
        'god_start': attest.god_start,
        'data_end': end_date_for_calendar,
        'data4_end': attest.data4_end,
        'god_end': attest.god_end,
        'vid': attest.vid,
        'attest_kod': attest.kod,
        'mesto': attest.mesto,
        'attest_adress': attest.adress,
        'ruka': attest.ruka,
        'num_dogovor': contract.num_dogovor,
        'profORG': contract.profORG,
        'ryco': contract.ryco,
        'contract_adress': contract.adress,
        'smart': contract.smart,
        'rec': contract.rec,
        'control_gru': control.gru,
        'control_code': control.code,
        'napra': control.napra,
        'control_data': control.data,
        'rycSPO': control.rycSPO,
        'date_check': control.date_check,
        'pred': control.pred,
        'rycP': control.rycP,
        'control_fio': control.fio,
        'fioO': control.fioO,
        'pri': control.pri,
        'field1': type_practic.field1,
        'appendix_addres': appendix.addres,
        'namePomech': appendix.namePomech,
    }
    
    return render(request, "profile.html", context)


def logout_view(request):
    logout(request)
    messages.success(request, 'Вы вышли из системы')
    return redirect('/login/')


def download_blank_templates(request):
    """
    Скачивание всех пустых шаблонов из папки blank_templates
    """
    # Путь к папке с пустыми шаблонами
    blanks_dir = os.path.join(settings.BASE_DIR, 'templates', 'blank_templates')
    
    # Для отладки - выводим в консоль
    print(f"BASE_DIR: {settings.BASE_DIR}")
    print(f"Ищем шаблоны в: {blanks_dir}")
    print(f"Папка существует: {os.path.exists(blanks_dir)}")
    
    if not os.path.exists(blanks_dir):
        messages.error(request, f'❌ Папка с шаблонами не найдена: {blanks_dir}')
        return redirect('/profile/')
    
    # Получаем все .docx файлы из папки
    blank_files = []
    for file in os.listdir(blanks_dir):
        if file.endswith('.docx'):
            blank_files.append(file)
            print(f"Найден файл: {file}")
    
    if not blank_files:
        messages.error(request, f'❌ В папке {blanks_dir} нет файлов .docx')
        return redirect('/profile/')
    
    zip_buffer = BytesIO()
    downloaded_count = 0
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for filename in blank_files:
            file_path = os.path.join(blanks_dir, filename)
            if os.path.exists(file_path):
                zip_file.write(file_path, filename)
                downloaded_count += 1
                print(f"✓ Добавлен в архив: {filename}")
    
    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename="Пустые_шаблоны_документов.zip"'
    messages.success(request, f'✅ Скачано {downloaded_count} пустых шаблонов!')
    return response